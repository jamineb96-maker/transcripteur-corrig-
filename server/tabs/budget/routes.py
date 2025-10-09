"""Routes API pour l'onglet Budget cognitif et somatique."""

from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List

from flask import Response, current_app, jsonify, request, send_from_directory, url_for

try:  # pragma: no cover - dépendance optionnelle
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfgen import canvas
    REPORTLAB_AVAILABLE = True
except ModuleNotFoundError:  # pragma: no cover - fallback pour les tests
    A4 = (595.27, 841.89)  # type: ignore[assignment]
    mm = 2.83465  # type: ignore[assignment]
    canvas = None  # type: ignore[assignment]
    REPORTLAB_AVAILABLE = False

from server.services import graphs
from server.services.budget_engine import (
    AssessmentResult,
    BudgetComputationError,
    PatientContext,
    compute_assessment,
    export_basename,
    list_presets,
    result_from_dict,
    summarize_for_history,
)
from server.services.budget_engine import update_profile_bias
from . import bp


def _history_root() -> Path:
    base = Path(current_app.instance_path)
    path = base / 'budget_history'
    path.mkdir(parents=True, exist_ok=True)
    return path


def _history_file(patient_id: str) -> Path:
    return _history_root() / f"{patient_id}.json"


def _exports_dir() -> Path:
    path = _history_root() / 'exports'
    path.mkdir(parents=True, exist_ok=True)
    return path


def _load_record(patient_id: str) -> Dict[str, Any]:
    path = _history_file(patient_id)
    if not path.exists():
        return {}
    try:
        with path.open('r', encoding='utf-8') as handle:
            return json.load(handle)
    except Exception:
        current_app.logger.warning('Unable to read budget record for %s', patient_id)
        return {}


def _save_record(patient_id: str, data: Dict[str, Any]) -> None:
    path = _history_file(patient_id)
    with path.open('w', encoding='utf-8') as handle:
        json.dump(data, handle, ensure_ascii=False, indent=2)


def _json_error(status: int, code: str, message: str) -> Response:
    payload = {'success': False, 'error': {'code': code, 'message': message}}
    response = jsonify(payload)
    response.status_code = status
    return response


def _generate_pdf(path: Path, patient: PatientContext, result: AssessmentResult) -> Path:
    if not REPORTLAB_AVAILABLE:
        raise RuntimeError('reportlab_missing')
    pdf = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    margin = 20 * mm

    pdf.setTitle('Budget cognitif et somatique')
    pdf.setAuthor('Assistant clinique')

    # Page de garde
    pdf.setFont('Helvetica-Bold', 20)
    pdf.drawString(margin, height - margin - 20, 'Budget cognitif et somatique')
    pdf.setFont('Helvetica', 12)
    pdf.drawString(margin, height - margin - 60, f"Patient : {patient.name}")
    pdf.drawString(margin, height - margin - 80, f"Période : {'Journée' if patient.period == 'day' else 'Semaine'}")
    pdf.drawString(margin, height - margin - 100, f"Profil : {patient.budget_profile}")
    pdf.drawString(margin, height - margin - 120, datetime.utcnow().strftime('Édition du %d/%m/%Y à %Hh%M'))
    pdf.showPage()

    # Page estimation
    pdf.setFont('Helvetica-Bold', 16)
    pdf.drawString(margin, height - margin - 20, 'Estimation et hypothèses')
    pdf.setFont('Helvetica', 11)
    text = pdf.beginText(margin, height - margin - 50)
    text.setLeading(14)
    text.textLine(f"Stock estimé : {result.spoons_stock:.1f} cuillères")
    cost = sum(item.value for item in result.spoons_consumption)
    recovery = sum(item.value for item in result.spoons_recovery)
    net = result.net_spoons_day
    text.textLine(f"Coût cumulé : {cost:.1f} cuillères")
    text.textLine(f"Récupération potentielle : {recovery:.1f} cuillères")
    text.textLine(f"Solde projeté : {net:.1f} cuillères")
    pdf.drawText(text)

    balance = graphs.build_spoon_balance(result.spoons_stock, cost, recovery)
    graphs.draw_on_canvas(pdf, balance, margin, height / 2 - 40)
    timeline = graphs.build_timeline(
        result.spoons_consumption,
        result.spoons_recovery,
        'journée' if patient.period == 'day' else 'semaine',
    )
    graphs.draw_on_canvas(pdf, timeline, margin, margin + 20)
    pdf.showPage()

    # Page psychoéducation
    pdf.setFont('Helvetica-Bold', 16)
    pdf.drawString(margin, height - margin - 20, 'Psychoéducation critique')
    pdf.setFont('Helvetica', 11)
    text = pdf.beginText(margin, height - margin - 50)
    text.setLeading(15)
    for paragraph in result.narrative_summary.split('\n'):
        text.textLines(paragraph.strip())
        text.textLine('')
    if result.alerts:
        for alert in result.alerts:
            text.textLine(f"Alerte : {alert['message']}")
    pdf.drawText(text)
    pdf.save()

    if path.stat().st_size < 6_000:
        raise RuntimeError('pdf_too_small')
    return path


def _generate_docx(path: Path, patient: PatientContext, result: AssessmentResult) -> Path:
    try:
        from docx import Document
        from docx.shared import Inches
    except ModuleNotFoundError as exc:  # pragma: no cover - dépendance optionnelle
        raise RuntimeError('docx_dependency_missing') from exc

    document = Document()
    document.add_heading('Budget cognitif et somatique', level=1)
    document.add_paragraph(f"Patient : {patient.name}")
    document.add_paragraph(f"Période : {'Journée' if patient.period == 'day' else 'Semaine'}")
    document.add_paragraph(f"Profil : {patient.budget_profile}")

    cost = sum(item.value for item in result.spoons_consumption)
    recovery = sum(item.value for item in result.spoons_recovery)

    document.add_heading('Estimation', level=2)
    document.add_paragraph(f"Stock estimé : {result.spoons_stock:.1f} cuillères")
    document.add_paragraph(f"Coût cumulé : {cost:.1f} cuillères")
    document.add_paragraph(f"Récupération potentielle : {recovery:.1f} cuillères")
    document.add_paragraph(f"Solde projeté : {result.net_spoons_day:.1f} cuillères")

    try:
        balance = graphs.build_spoon_balance(result.spoons_stock, cost, recovery)
        timeline = graphs.build_timeline(
            result.spoons_consumption,
            result.spoons_recovery,
            'journée' if patient.period == 'day' else 'semaine',
        )
        balance_path = graphs.save_as_png(balance, path.with_suffix('.balance.png'))
        timeline_path = graphs.save_as_png(timeline, path.with_suffix('.timeline.png'))
    except RuntimeError as exc:
        if str(exc) == 'reportlab_missing':
            raise RuntimeError('reportlab_missing') from exc
        raise
    document.add_picture(str(balance_path), width=Inches(5.5))
    document.add_picture(str(timeline_path), width=Inches(5.5))

    document.add_heading('Psychoéducation critique', level=2)
    document.add_paragraph(result.narrative_summary)
    if result.alerts:
        for alert in result.alerts:
            document.add_paragraph(alert['message'])

    document.save(str(path))
    if path.stat().st_size < 6_000:
        raise RuntimeError('docx_too_small')
    return path


@bp.get('/presets')
def presets() -> Response:
    return jsonify({'success': True, 'data': list_presets()})


@bp.post('/assess')
def assess() -> Response:
    payload = request.get_json(silent=True) or {}
    try:
        patient = PatientContext.from_payload(payload.get('patient', {}))
    except BudgetComputationError as exc:
        return _json_error(400, str(exc), 'Paramètres patient invalides.')

    record = _load_record(patient.id)
    profile = record.get('profile', {})

    try:
        result = compute_assessment(payload, profile)
    except BudgetComputationError as exc:
        return _json_error(400, str(exc), 'Évaluation impossible.')
    except FileNotFoundError:
        return _json_error(500, 'missing_presets', 'Le fichier de presets est manquant.')

    response_data = result.to_dict()

    if payload.get('persist'):
        history_entry = {
            'timestamp': datetime.utcnow().isoformat(timespec='seconds') + 'Z',
            'summary': summarize_for_history(result),
            'result': response_data,
        }
        history = record.setdefault('history', [])
        history.append(history_entry)
        record['profile'] = profile
        record['patient'] = {
            'id': patient.id,
            'name': patient.name,
            'gender': patient.gender,
            'language': patient.language,
        }
        _save_record(patient.id, record)

    return jsonify(
        {
            'success': True,
            'data': response_data,
            'meta': {
                'profile_bias': profile.get('profile_bias', {}),
                'history_length': len(record.get('history', [])),
            },
        }
    )


@bp.post('/export')
def export_budget() -> Response:
    payload = request.get_json(silent=True) or {}
    try:
        patient = PatientContext.from_payload(payload.get('patient', {}))
    except BudgetComputationError as exc:
        return _json_error(400, str(exc), 'Patient invalide pour export.')

    result_data = payload.get('result')
    if not isinstance(result_data, dict):
        return _json_error(400, 'missing_result', "Résultat d'évaluation absent pour l'export.")

    result = result_from_dict(result_data)

    formats: List[str] = payload.get('formats') or ['pdf']
    basename = export_basename(patient) + '_' + datetime.utcnow().strftime('%Y%m%d_%H%M%S')
    export_dir = _exports_dir()

    files: Dict[str, str] = {}
    warnings: List[str] = []

    if 'pdf' in formats:
        try:
            pdf_path = _generate_pdf(export_dir / f"{basename}.pdf", patient, result)
            files['pdf'] = url_for('budget.download_export', filename=pdf_path.name)
        except RuntimeError as exc:
            if str(exc) == 'reportlab_missing':
                warnings.append('Export PDF indisponible (ReportLab non installé).')
            else:
                return _json_error(500, str(exc), 'Impossible de générer le PDF.')

    if 'docx' in formats:
        try:
            docx_path = _generate_docx(export_dir / f"{basename}.docx", patient, result)
            files['docx'] = url_for('budget.download_export', filename=docx_path.name)
        except RuntimeError as exc:
            if str(exc) == 'reportlab_missing':
                warnings.append('Export DOCX indisponible (ReportLab non installé).')
            else:
                warnings.append("Export DOCX indisponible (dépendance manquante ou taille insuffisante).")

    if not files:
        return _json_error(400, 'no_export', 'Aucun export disponible.')

    return jsonify({'success': True, 'data': files, 'warnings': warnings})


@bp.post('/save-profile')
def save_profile() -> Response:
    payload = request.get_json(silent=True) or {}
    patient_id = str(payload.get('patient_id') or '').strip()
    if not patient_id:
        return _json_error(400, 'missing_patient', 'Identifiant patient manquant.')

    record = _load_record(patient_id)
    profile = record.get('profile', {})
    category = str(payload.get('category') or '').strip()
    if not category:
        return _json_error(400, 'missing_category', 'Catégorie à calibrer manquante.')
    try:
        delta = float(payload.get('delta', 0.0))
    except (TypeError, ValueError):
        return _json_error(400, 'invalid_delta', 'Delta de calibration invalide.')

    profile = update_profile_bias(profile, category, delta)
    record['profile'] = profile
    _save_record(patient_id, record)

    return jsonify({'success': True, 'data': profile})


@bp.get('/history')
def history() -> Response:
    patient_id = request.args.get('patient')
    if not patient_id:
        return _json_error(400, 'missing_patient', 'Identifiant patient requis.')
    record = _load_record(patient_id)
    history_entries = record.get('history', [])
    return jsonify({'success': True, 'data': history_entries})


@bp.get('/download/<path:filename>')
def download_export(filename: str):
    directory = _exports_dir()
    target = directory / filename
    if not target.exists():
        return _json_error(404, 'missing_file', 'Fichier introuvable.')
    return send_from_directory(directory, filename, as_attachment=True)
