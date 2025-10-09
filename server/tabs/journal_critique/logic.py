"""Logique métier pour l'onglet Journal critique."""

from __future__ import annotations

import base64
import io
import json
import re
import textwrap
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Sequence, Tuple

from server.services.paths import ensure_patient_subdir
from server.util import slugify

try:  # pragma: no cover - optional dependency evaluated at runtime
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer
    REPORTLAB_AVAILABLE = True
except ModuleNotFoundError:  # pragma: no cover
    colors = None
    ParagraphStyle = None
    getSampleStyleSheet = None
    SimpleDocTemplate = None
    PageBreak = None
    Paragraph = None
    Spacer = None
    A4 = None
    mm = 1
    REPORTLAB_AVAILABLE = False

try:  # pragma: no cover - optional dependency evaluated at runtime
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ModuleNotFoundError:  # pragma: no cover
    Document = None
    WD_STYLE_TYPE = None
    WD_ALIGN_PARAGRAPH = None
    Pt = None
    DOCX_AVAILABLE = False


PROMPTS_INDEX_PATH = Path("library/journal_prompts_index.json")
PROMPTS_BASE_DIR = Path("library/journal_prompts")

# Domaines pour l'évaluation de couverture
DOMAINS = ["somatique", "cognitif", "relationnel", "politique", "valeurs"]

# Mots bannis pour les validations éditoriales
BANNED_PATTERNS = [
    r"psychanal",
    r"borderline",
    r"trouble de la personnalité",
    r"non[ -]?compliant",
    r"observance",
    r"protocole rigide",
    r"normalisation comportementale",
]

FAMILIES_LABELS = {
    "externalisation": "Externalisation du problème",
    "resultats_uniques": "Résultats uniques",
    "re_membering": "Re-membering",
    "dialogues_internes": "Dialogues internes situés",
    "cartographies": "Cartographies somato-cognitives",
    "relationnel": "Écritures relationnelles",
    "documents": "Lettres et documents narratifs",
}

# Correspondances domaines ↔ familles
FAMILY_DOMAINS = {
    "externalisation": {"politique", "relationnel"},
    "resultats_uniques": {"valeurs", "relationnel"},
    "re_membering": {"relationnel", "valeurs"},
    "dialogues_internes": {"cognitif", "valeurs"},
    "cartographies": {"somatique", "cognitif"},
    "relationnel": {"relationnel", "politique"},
    "documents": {"valeurs", "politique"},
}

# Mapping lenses → familles prioritaires
LENS_TO_FAMILIES = {
    "validisme": ["cartographies", "externalisation"],
    "patriarcat": ["relationnel", "documents"],
    "racisme": ["relationnel", "documents"],
    "classisme": ["relationnel", "documents"],
    "cisheteronormativite": ["relationnel", "externalisation"],
    "colonialite": ["externalisation", "documents"],
}

_TEMPLATE_BASE = """# {{TITLE}}
{{INTRO: contexte bref, situer le problème hors de la personne, nommer les rapports de pouvoir pertinents.}}

{{INVITATION_PRINCIPALE}}
{{TUT|VOUS}} peux/pouvez explorer comment {{GEN:F|M|N}} as/avez déjà, même légèrement, contredit la stratégie du Problème dans une situation précise. Écris/Écrivez ce qui s’est rendu possible, qui l’a vu, et ce que cela raconte de {{GEN:tes|vos|tes}} valeurs.

{{VARIANTE_BUDGET_FAIBLE}}
Si l’énergie est basse, choisis/choisissez un seul moment de 2–3 lignes et note/Notez uniquement le détail le plus étonnant qui n’allait pas dans le sens du Problème.

{{ENCADRE_CONTEXTUALISATION}}
Relie/Reliez cet épisode aux conditions matérielles et sociales (temps, argent, accès, normes) qui le rendent plus ou moins possible.

{{TEMOIN_OUTSIDER_WITNESS}}
Qui pourrait en témoigner dans {{GEN:ton|votre|ton}} entourage et quels mots utiliserait-il pour décrire ce qu’il voit de {{GEN:toi|vous|toi}} quand le Problème perd un peu de terrain ?
"""

RECOMMENDATION_TEMPLATES = {
    "somatique": {
        "title": "Gabarits somatiques situés",
        "suggestions": [
            {
                "title": "Journal corporel non performant",
                "tags": ["validisme", "neurodiversite", "handicaps"],
                "template": _TEMPLATE_BASE,
            },
        ],
    },
    "cognitif": {
        "title": "Cartographies attentionnelles",
        "suggestions": [
            {
                "title": "Pacte attentionnel protecteur",
                "tags": ["neurodiversite", "anticulpabilite"],
                "template": _TEMPLATE_BASE,
            },
        ],
    },
    "relationnel": {
        "title": "Alliances et justice relationnelle",
        "suggestions": [
            {
                "title": "Lettre de coalition",
                "tags": ["justice", "alliances"],
                "template": _TEMPLATE_BASE,
            },
        ],
    },
    "politique": {
        "title": "Contextualisations politiques",
        "suggestions": [
            {
                "title": "Déclaration anti-stigmate",
                "tags": ["validisme", "patriarcat", "racisme"],
                "template": _TEMPLATE_BASE,
            },
        ],
    },
    "valeurs": {
        "title": "Résultats uniques à amplifier",
        "suggestions": [
            {
                "title": "Certificat de pratiques situées",
                "tags": ["re_authoring", "documentation"],
                "template": _TEMPLATE_BASE,
            },
        ],
    },
}


@dataclass
class Prompt:
    id: str
    title: str
    family: str
    tags: List[str]
    reading_level: str
    budget_profile: str
    contraindications: List[str]
    md_file: Path
    domains: List[str] = field(default_factory=list)

    def to_dict(self) -> Dict[str, object]:
        return {
            "id": self.id,
            "title": self.title,
            "family": self.family,
            "familyLabel": FAMILIES_LABELS.get(self.family, self.family.title()),
            "tags": self.tags,
            "reading_level": self.reading_level,
            "budget_profile": self.budget_profile,
            "contraindications": self.contraindications,
            "domains": self.domains,
        }


_PROMPT_CACHE: Dict[str, Prompt] = {}


def load_prompts() -> Dict[str, Prompt]:
    """Charge l'index des prompts avec mise en cache."""

    global _PROMPT_CACHE
    if _PROMPT_CACHE:
        return _PROMPT_CACHE
    if not PROMPTS_INDEX_PATH.exists():
        raise FileNotFoundError("journal_prompts_index.json manquant")
    data = json.loads(PROMPTS_INDEX_PATH.read_text(encoding="utf-8"))
    cache: Dict[str, Prompt] = {}
    for entry in data:
        md_path = PROMPTS_BASE_DIR / entry["md_file"].split("/")[-1]
        prompt = Prompt(
            id=entry["id"],
            title=entry["title"],
            family=entry["family"],
            tags=entry.get("tags", []),
            reading_level=entry.get("reading_level", "intermediaire"),
            budget_profile=entry.get("budget_profile", "moyen"),
            contraindications=entry.get("contraindications", []),
            md_file=md_path,
            domains=entry.get("domains", list(FAMILY_DOMAINS.get(entry.get("family"), []))),
        )
        cache[prompt.id] = prompt
    _PROMPT_CACHE = cache
    return cache


def list_prompts(filters: Optional[Dict[str, str]] = None) -> List[Dict[str, object]]:
    prompts = load_prompts()
    result: List[Dict[str, object]] = []
    filters = filters or {}
    query = (filters.get("q") or filters.get("query") or "").strip().lower()
    family = (filters.get("family") or "").strip().lower()
    tag_filter = set(filter(None, (filters.get("tags") or "").lower().split(",")))
    budget = (filters.get("budget") or "").strip().lower()
    reading = (filters.get("reading_level") or "").strip().lower()
    lens = (filters.get("lens") or "").strip().lower()

    for prompt in prompts.values():
        if family and prompt.family != family:
            continue
        if budget and prompt.budget_profile != budget:
            continue
        if reading and prompt.reading_level != reading:
            continue
        if tag_filter and not (tag_filter & set(t.lower() for t in prompt.tags)):
            continue
        if lens and lens not in (t.lower() for t in prompt.tags):
            continue
        if query:
            haystack = " ".join([prompt.title, prompt.family, " ".join(prompt.tags)])
            if query not in haystack.lower():
                try:
                    text = prompt.md_file.read_text(encoding="utf-8")
                except FileNotFoundError:
                    text = ""
                if query not in text.lower():
                    continue
        result.append(prompt.to_dict())
    result.sort(key=lambda item: (item["family"], item["title"]))
    return result


def get_prompt_content(prompt_id: str) -> Tuple[Prompt, str]:
    prompts = load_prompts()
    if prompt_id not in prompts:
        raise KeyError(prompt_id)
    prompt = prompts[prompt_id]
    content = prompt.md_file.read_text(encoding="utf-8")
    return prompt, content


def _replace_tokens(text: str, *, langage: str, gender: str, patient: Dict[str, str], tempo: str) -> str:
    """Applique les tokens de civilité et de genre."""

    langage = langage or "tu"
    gender = gender or "neutral"
    tempo = tempo or "present"

    def repl(match):
        raw = match.group(1)
        upper = raw.upper()
        if upper.startswith("TUT:"):
            values = raw.split(":", 1)[1].split("|")
            first = values[0] if values else ""
            second = values[1] if len(values) > 1 else first
            return first if langage == "tu" else second
        if upper == "TUT|VOUS":
            return "Tu" if langage == "tu" else "Vous"
        if upper.startswith("GEN:"):
            values = raw.split(":", 1)[1]
            options = values.split("|")
            mapping = {
                "feminine": options[0] if len(options) > 0 else options[-1],
                "masculine": options[1] if len(options) > 1 else options[0],
                "neutral": options[2] if len(options) > 2 else options[-1],
            }
            return mapping.get(gender, mapping["neutral"])
        if upper.startswith("TEMPS:"):
            values = raw.split(":", 1)[1].split("|")
            mapping = {
                "present": values[0] if len(values) > 0 else values[-1],
                "futur": values[1] if len(values) > 1 else values[0],
            }
            return mapping.get(tempo, mapping["present"])
        if upper == "PATIENT_PRENOM":
            return patient.get("name") or "la personne"
        if upper == "PRONOM_SUJET":
            return "tu" if langage == "tu" else "vous"
        if upper == "PRONOM_OBJET":
            return "te" if langage == "tu" else "vous"
        if upper == "PRONOM_POSSESSIF":
            return {"tu": "ton", "vous": "votre"}.get(langage, "ton")
        return raw

    pattern = re.compile(r"\{\{([^{}]+)\}\}")
    text = pattern.sub(repl, text)
    return text


@dataclass
class PromptRender:
    prompt: Prompt
    title: str
    intro: List[str]
    invitation: List[str]
    variante: List[str]
    contextualisation: List[str]
    witness: List[str]


def parse_prompt_markdown(prompt: Prompt, content: str) -> PromptRender:
    lines = content.splitlines()
    title = prompt.title
    sections = {
        "intro": [],
        "invitation": [],
        "variante": [],
        "contextualisation": [],
        "witness": [],
    }
    current = "intro"
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("# "):
            title = stripped[2:].strip() or prompt.title
            continue
        if stripped.lower().startswith("## invitation"):
            current = "invitation"
            continue
        if stripped.lower().startswith("## variante"):
            current = "variante"
            continue
        if stripped.lower().startswith("## encadr"):
            current = "contextualisation"
            continue
        if "temoin" in stripped.lower():
            current = "witness"
            continue
        if stripped:
            sections[current].append(stripped)
    return PromptRender(
        prompt=prompt,
        title=title,
        intro=sections["intro"],
        invitation=sections["invitation"],
        variante=sections["variante"],
        contextualisation=sections["contextualisation"],
        witness=sections["witness"],
    )


def render_prompt(prompt_id: str, *, langage: str, gender: str, patient: Dict[str, str], tempo: str) -> PromptRender:
    try:
        prompt, content = get_prompt_content(prompt_id)
    except KeyError as exc:
        raise ValueError("unknown_prompt") from exc
    except FileNotFoundError as exc:
        raise ValueError("missing_prompt_file") from exc
    content = _replace_tokens(content, langage=langage, gender=gender, patient=patient, tempo=tempo)
    return parse_prompt_markdown(prompt, content)


def check_prohibited_language(rendered_prompts: Sequence[PromptRender]) -> None:
    for render in rendered_prompts:
        joined = "\n".join(
            render.intro + render.invitation + render.variante + render.contextualisation + render.witness
        )
        for pattern in BANNED_PATTERNS:
            if re.search(pattern, joined, flags=re.IGNORECASE):
                raise ValueError(
                    "prompt_banned_language",
                )


def ensure_required_families(rendered_prompts: Sequence[PromptRender]) -> None:
    families = {render.prompt.family for render in rendered_prompts}
    if "externalisation" not in families:
        raise ValueError("missing_externalisation")
    if "resultats_uniques" not in families:
        raise ValueError("missing_resultats_uniques")


def _build_cover(patient: Dict[str, str], langage: str, gender: str) -> Dict[str, str]:
    civility = "Tutoiement" if langage == "tu" else "Vouvoiement"
    gender_label = {
        "feminine": "Genre grammatical : féminin",
        "masculine": "Genre grammatical : masculin",
        "neutral": "Genre grammatical : neutre",
    }.get(gender, "Genre grammatical : neutre")
    return {
        "title": "Journal critique",
        "patient": patient.get("name", "Patient·e"),
        "date": datetime.now().strftime("%d/%m/%Y"),
        "civility": civility,
        "gender": gender_label,
        "reminder": (
            "Document narratif non prescriptif. Les invitations ci-dessous"
            " soutiennent la re-authorisation face aux stigmates identifiés."
        ),
    }


def _build_annexes(artefacts: Dict[str, object], langage: str) -> List[Tuple[str, List[str]]]:
    annexes: List[Tuple[str, List[str]]] = []
    evidence = artefacts.get("evidence_sheet")
    if isinstance(evidence, str) and evidence.strip():
        paragraphs = [
            "Synthèse située des matériaux mobilisés lors de la séance,"
            " à revisiter librement pour nourrir les écritures.",
        ]
        paragraphs.extend(textwrap.wrap(evidence.strip(), 420))
        annexes.append(("Bibliographie située", paragraphs))
    if artefacts.get("reperes_candidates"):
        paragraphs = [
            "Mini-lettre outsider-witness : ce document condense une trace"
            " de témoignage solidaire à mobiliser lorsque cela soutient la personne.",
        ]
        for idx, item in enumerate(artefacts["reperes_candidates"][:2]):
            if not isinstance(item, dict):
                continue
            title = item.get("title", f"Témoin {idx+1}")
            body = item.get("body", "")
            paragraphs.append(f"{title} — {body}")
        annexes.append(("Lettre outsider-witness", paragraphs))
    return annexes


def _pdf_styles() -> Dict[str, ParagraphStyle]:
    if not REPORTLAB_AVAILABLE:  # pragma: no cover - handled upstream
        raise RuntimeError("reportlab_missing")
    styles = getSampleStyleSheet()
    base = styles["BodyText"]
    styles.add(
        ParagraphStyle(
            "PromptTitle",
            parent=styles["Heading2"],
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            "FamilyHeading",
            parent=styles["Heading1"],
            spaceAfter=12,
        )
    )
    styles.add(
        ParagraphStyle(
            "ContextBox",
            parent=base,
            backColor=colors.whitesmoke,
            borderColor=colors.lightgrey,
            borderWidth=0.5,
            borderPadding=6,
            spaceBefore=6,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            "AnnexeHeading",
            parent=styles["Heading2"],
            spaceBefore=18,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            "CoverTitle",
            parent=styles["Heading1"],
            alignment=1,
            fontSize=24,
            spaceAfter=20,
        )
    )
    styles.add(
        ParagraphStyle(
            "CoverMeta",
            parent=base,
            alignment=1,
            leading=16,
        )
    )
    return styles


def build_pdf(rendered_prompts: Sequence[PromptRender], *, langage: str, gender: str, patient: Dict[str, str], artefacts: Dict[str, object]) -> bytes:
    if not REPORTLAB_AVAILABLE:
        raise ValueError("reportlab_missing")
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
    )
    styles = _pdf_styles()
    story: List[object] = []

    cover = _build_cover(patient, langage, gender)
    story.append(Paragraph(cover["title"], styles["CoverTitle"]))
    story.append(Paragraph(f"Patient·e : {cover['patient']}", styles["CoverMeta"]))
    story.append(Paragraph(f"Date : {cover['date']}", styles["CoverMeta"]))
    story.append(Paragraph(cover["civility"], styles["CoverMeta"]))
    story.append(Paragraph(cover["gender"], styles["CoverMeta"]))
    story.append(Spacer(1, 18))
    story.append(Paragraph(cover["reminder"], styles["BodyText"]))
    story.append(PageBreak())

    grouped: Dict[str, List[PromptRender]] = {}
    for render in rendered_prompts:
        grouped.setdefault(render.prompt.family, []).append(render)
    grouped_items = list(grouped.items())
    annexes = _build_annexes(artefacts, langage)

    for index, (family, items) in enumerate(grouped_items):
        story.append(Paragraph(FAMILIES_LABELS.get(family, family.title()), styles["FamilyHeading"]))
        for render in items:
            story.append(Paragraph(render.title, styles["PromptTitle"]))
            for block in render.intro:
                story.append(Paragraph(block, styles["BodyText"]))
            if render.invitation:
                story.append(Paragraph("Invitation principale", styles["Heading3"]))
                for block in render.invitation:
                    story.append(Paragraph(block, styles["BodyText"]))
            if render.variante:
                story.append(Paragraph("Variante budget faible", styles["Heading3"]))
                for block in render.variante:
                    story.append(Paragraph(block, styles["BodyText"]))
            if render.contextualisation:
                story.append(Paragraph("Encadré – situer le contexte", styles["Heading3"]))
                for block in render.contextualisation:
                    story.append(Paragraph(block, styles["ContextBox"]))
            if render.witness:
                story.append(Paragraph("Témoin outsider-witness", styles["Heading3"]))
                for block in render.witness:
                    story.append(Paragraph(block, styles["BodyText"]))
            story.append(Spacer(1, 12))
        if index < len(grouped_items) - 1 or annexes:
            story.append(PageBreak())

    if annexes:
        story.append(Paragraph("Annexes", styles["FamilyHeading"]))
        for title, paragraphs in annexes:
            story.append(Paragraph(title, styles["AnnexeHeading"]))
            for block in paragraphs:
                story.append(Paragraph(block, styles["BodyText"]))
            story.append(Spacer(1, 12))

    doc.build(story)
    return buffer.getvalue()


def _docx_styles(document: Document) -> None:
    if not DOCX_AVAILABLE:
        raise RuntimeError("docx_missing")
    styles = document.styles
    if "FamilyHeading" not in styles:
        style = styles.add_style("FamilyHeading", WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(18)
        style.font.bold = True
    if "PromptTitle" not in styles:
        style = styles.add_style("PromptTitle", WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(14)
        style.font.bold = True
    if "ContextBox" not in styles:
        style = styles.add_style("ContextBox", WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(11)
        style.font.italic = True


def build_docx(rendered_prompts: Sequence[PromptRender], *, langage: str, gender: str, patient: Dict[str, str], artefacts: Dict[str, object]) -> bytes:
    if not DOCX_AVAILABLE:
        raise ValueError("docx_missing")
    document = Document()
    _docx_styles(document)
    cover = _build_cover(patient, langage, gender)
    title = document.add_heading(cover["title"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for key in ["patient", "date", "civility", "gender"]:
        para = document.add_paragraph(cover[key])
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(cover["reminder"])
    document.add_page_break()

    grouped: Dict[str, List[PromptRender]] = {}
    for render in rendered_prompts:
        grouped.setdefault(render.prompt.family, []).append(render)
    grouped_items = list(grouped.items())
    annexes = _build_annexes(artefacts, langage)

    for index, (family, items) in enumerate(grouped_items):
        document.add_paragraph(FAMILIES_LABELS.get(family, family.title()), style="FamilyHeading")
        for render in items:
            document.add_paragraph(render.title, style="PromptTitle")
            for block in render.intro:
                document.add_paragraph(block)
            if render.invitation:
                document.add_heading("Invitation principale", level=2)
                for block in render.invitation:
                    document.add_paragraph(block)
            if render.variante:
                document.add_heading("Variante budget faible", level=2)
                for block in render.variante:
                    document.add_paragraph(block)
            if render.contextualisation:
                document.add_heading("Encadré – situer le contexte", level=2)
                for block in render.contextualisation:
                    document.add_paragraph(block, style="ContextBox")
            if render.witness:
                document.add_heading("Témoin outsider-witness", level=2)
                for block in render.witness:
                    document.add_paragraph(block)
            document.add_paragraph("")
        if index < len(grouped_items) - 1 or annexes:
            document.add_page_break()

    if annexes:
        document.add_paragraph("Annexes", style="FamilyHeading")
        for title, paragraphs in annexes:
            document.add_paragraph(title, style="PromptTitle")
            for block in paragraphs:
                document.add_paragraph(block)
            document.add_paragraph("")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class JournalStorage:
    def _patient_dir(self, patient_id: str) -> Path:
        slug = slugify(patient_id or "")
        return ensure_patient_subdir(slug, "journal_critique")

    def save_document(
        self,
        *,
        patient_id: str,
        payload: Dict[str, object],
        pdf_bytes: bytes,
        docx_bytes: bytes,
        rendered_prompts: Sequence[PromptRender],
        artefacts: Dict[str, object],
    ) -> Dict[str, object]:
        timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        patient_dir = self._patient_dir(patient_id)
        pdf_path = patient_dir / f"{timestamp}_journal.pdf"
        docx_path = patient_dir / f"{timestamp}_journal.docx"
        pdf_path.write_bytes(pdf_bytes)
        docx_path.write_bytes(docx_bytes)
        slug = patient_dir.parent.name
        relative_pdf = f"{slug}/journal_critique/{pdf_path.name}"
        relative_docx = f"{slug}/journal_critique/{docx_path.name}"
        history_entry = {
            "patient_id": patient_id,
            "patient_name": payload.get("patient", {}).get("name"),
            "timestamp": timestamp,
            "pdf": relative_pdf,
            "docx": relative_docx,
            "selected_prompts": [render.prompt.id for render in rendered_prompts],
            "coverage": payload.get("coverage"),
            "alerts": payload.get("alerts"),
        }
        index_path = patient_dir / "history.json"
        history: List[Dict[str, object]] = []
        if index_path.exists():
            try:
                history = json.loads(index_path.read_text(encoding="utf-8"))
            except json.JSONDecodeError:
                history = []
        history.append(history_entry)
        index_path.write_text(json.dumps(history, ensure_ascii=False, indent=2), encoding="utf-8")
        return {
            "pdf_path": str(pdf_path),
            "docx_path": str(docx_path),
            "pdf_relative": relative_pdf,
            "docx_relative": relative_docx,
            "entry": history_entry,
        }

    def history(self, patient_id: Optional[str] = None) -> List[Dict[str, object]]:
        if patient_id:
            index_path = self._patient_dir(patient_id) / "history.json"
            if not index_path.exists():
                return []
            try:
                history = json.loads(index_path.read_text(encoding="utf-8"))
            except json.JSONDecodeError:
                return []
            history.sort(key=lambda entry: entry.get("timestamp", ""), reverse=True)
            return history

        # Agrège l'historique de tous les patients présents dans les archives.
        entries: List[Dict[str, object]] = []
        archives_root = PROMPTS_BASE_DIR.parent.parent / "instance" / "archives"
        if archives_root.exists():
            for patient_dir in archives_root.iterdir():
                if not patient_dir.is_dir():
                    continue
                history_path = patient_dir / "journal_critique" / "history.json"
                if not history_path.exists():
                    continue
                try:
                    data = json.loads(history_path.read_text(encoding="utf-8"))
                except json.JSONDecodeError:
                    continue
                if isinstance(data, list):
                    entries.extend(data)
        entries.sort(key=lambda entry: entry.get("timestamp", ""), reverse=True)
        return entries


def suggest_prompts_from_postsession(artefacts: Dict[str, object], *, budget: str = "moyen", limit: int = 5) -> List[Dict[str, object]]:
    prompts = load_prompts()
    ranked: List[Tuple[int, Prompt, str]] = []
    somatiques = artefacts.get("indices_somatiques") or []
    cognitifs = artefacts.get("indices_cognitifs") or []
    contradictions = artefacts.get("contradiction_spans") or artefacts.get("contradictions") or []
    lenses = artefacts.get("lenses_used") or []
    budget = budget or "moyen"

    lens_slugs = []
    for lens in lenses:
        if isinstance(lens, dict) and lens.get("slug"):
            lens_slugs.append(lens["slug"].lower())
        elif isinstance(lens, str):
            lens_slugs.append(lens.lower())

    for prompt in prompts.values():
        score = 0
        reasons: List[str] = []
        if prompt.budget_profile == budget:
            score += 2
            reasons.append("budget aligné")
        if somatiques and "somatique" in prompt.domains:
            score += 3
            reasons.append("indices somatiques relevés")
        if cognitifs and "cognitif" in prompt.domains:
            score += 2
            reasons.append("indices cognitifs présents")
        if contradictions and prompt.family == "resultats_uniques":
            score += 3
            reasons.append("contradictions à explorer")
        for lens in lens_slugs:
            families = LENS_TO_FAMILIES.get(lens, [])
            if prompt.family in families:
                score += 2
                reasons.append(f"lens {lens}")
        if prompt.family == "documents" and artefacts.get("reperes_candidates"):
            score += 1
            reasons.append("documentation attendue")
        if score:
            ranked.append((score, prompt, ", ".join(sorted(set(reasons)))))
    ranked.sort(key=lambda item: item[0], reverse=True)
    suggestions: List[Dict[str, object]] = []
    for score, prompt, reason in ranked[:limit]:
        suggestions.append({"prompt": prompt.to_dict(), "justification": reason, "score": score})
    return suggestions


def assess_prompt_coverage(artefacts: Dict[str, object], selected: Sequence[str]) -> Dict[str, object]:
    prompts = load_prompts()
    domain_weights = {domain: 0.0 for domain in DOMAINS}
    availability = {domain: 0 for domain in DOMAINS}
    for prompt in prompts.values():
        for domain in prompt.domains:
            availability[domain] += 1
    for prompt_id in selected:
        prompt = prompts.get(prompt_id)
        if not prompt:
            continue
        for domain in prompt.domains:
            domain_weights[domain] += 1
    requirements = {domain: 1.0 for domain in DOMAINS}
    if artefacts.get("indices_somatiques"):
        requirements["somatique"] = 2.0
    if artefacts.get("indices_cognitifs"):
        requirements["cognitif"] = 2.0
    if artefacts.get("lenses_used"):
        requirements["politique"] = 2.0
    if artefacts.get("reperes_candidates"):
        requirements["valeurs"] = 2.0

    scores: Dict[str, int] = {}
    alerts: List[str] = []
    for domain, value in domain_weights.items():
        requirement = requirements.get(domain, 1.0)
        raw_score = 100 if requirement == 0 else min(100, int((value / requirement) * 100))
        scores[domain] = raw_score
        if raw_score < 60:
            alerts.append(f"Couverture {domain} faible")
        elif raw_score < 80:
            alerts.append(f"Couverture {domain} à renforcer")
        if availability.get(domain, 0) <= 1:
            alerts.append(f"Bibliothèque pauvre pour le domaine {domain}")
    if not selected:
        alerts.append("Aucun prompt sélectionné")
    return {"scores": scores, "alerts": sorted(set(alerts))}


def validate_budget_constraints(selected: Sequence[PromptRender], budget: str) -> None:
    if budget == "faible":
        for render in selected:
            if not render.variante:
                raise ValueError("missing_low_budget_variant")


def generate_preview(payload: Dict[str, object]) -> Dict[str, object]:
    prompts = payload.get("selected_prompts") or []
    if not prompts:
        raise ValueError("no_prompts")
    langage = payload.get("langage", "tu")
    gender = payload.get("genre", "neutral")
    patient = payload.get("patient", {})
    tempo = payload.get("tempo", "present")
    artefacts = payload.get("artefacts") or {}

    rendered = [
        render_prompt(prompt_id, langage=langage, gender=gender, patient=patient, tempo=tempo)
        for prompt_id in prompts
    ]
    check_prohibited_language(rendered)
    ensure_required_families(rendered)
    validate_budget_constraints(rendered, payload.get("budget_profile", "moyen"))
    pdf_bytes = build_pdf(rendered, langage=langage, gender=gender, patient=patient, artefacts=artefacts)
    if len(pdf_bytes) < 5 * 1024:
        raise ValueError("pdf_too_small")
    encoded = base64.b64encode(pdf_bytes).decode("ascii")
    return {"preview_pdf_base64": encoded}


def generate_document(payload: Dict[str, object]) -> Dict[str, object]:
    prompts = payload.get("selected_prompts") or []
    if not prompts:
        raise ValueError("no_prompts")
    langage = payload.get("langage", "tu")
    gender = payload.get("genre", "neutral")
    patient = payload.get("patient", {})
    tempo = payload.get("tempo", "present")
    artefacts = payload.get("artefacts") or {}

    rendered = [
        render_prompt(prompt_id, langage=langage, gender=gender, patient=patient, tempo=tempo)
        for prompt_id in prompts
    ]
    check_prohibited_language(rendered)
    ensure_required_families(rendered)
    validate_budget_constraints(rendered, payload.get("budget_profile", "moyen"))
    coverage = assess_prompt_coverage(artefacts, prompts)

    pdf_bytes = build_pdf(rendered, langage=langage, gender=gender, patient=patient, artefacts=artefacts)
    docx_bytes = build_docx(rendered, langage=langage, gender=gender, patient=patient, artefacts=artefacts)
    if len(pdf_bytes) < 5 * 1024:
        raise ValueError("pdf_too_small")
    if len(docx_bytes) < 5 * 1024:
        raise ValueError("docx_too_small")

    storage = JournalStorage()
    save_result = storage.save_document(
        patient_id=str(patient.get("id") or patient.get("name") or "patient"),
        payload={"patient": patient, "coverage": coverage, "alerts": coverage["alerts"]},
        pdf_bytes=pdf_bytes,
        docx_bytes=docx_bytes,
        rendered_prompts=rendered,
        artefacts=artefacts,
    )
    pdf_rel = save_result["entry"]["pdf"]
    docx_rel = save_result["entry"]["docx"]

    return {
        "pdf_url": f"/api/journal-critique/exports/{pdf_rel}",
        "docx_url": f"/api/journal-critique/exports/{docx_rel}",
        "history_entry": save_result["entry"],
        "coverage": coverage,
    }


def list_history(patient_id: Optional[str] = None) -> List[Dict[str, object]]:
    storage = JournalStorage()
    return storage.history(patient_id)


def get_recommendations(domain: str) -> Dict[str, object]:
    domain = (domain or "").lower()
    data = RECOMMENDATION_TEMPLATES.get(domain)
    if not data:
        return {"title": "", "suggestions": []}
    return data
