"""Utilities to export a simplified DOCX document from HTML fragments."""
from __future__ import annotations

import logging
from html.parser import HTMLParser
from pathlib import Path
from typing import List

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

LOGGER = logging.getLogger(__name__)


class _DocxHTMLParser(HTMLParser):
    def __init__(self, document: Document) -> None:
        super().__init__()
        self.document = document
        self._current_paragraph = document.add_paragraph()
        self._current_paragraph.style = document.styles["Normal"]
        self._current_runs: List = []
        self._list_level = 0

    # HTMLParser hooks -------------------------------------------------------
    def handle_starttag(self, tag: str, attrs):  # noqa: D401 - HTMLParser signature
        if tag in {"p", "div", "section"}:
            self._new_paragraph()
        elif tag in {"strong", "b"}:
            self._start_run(bold=True)
        elif tag in {"em", "i"}:
            self._start_run(italic=True)
        elif tag in {"ul", "ol"}:
            self._list_level += 1
        elif tag == "li":
            self._new_paragraph(list_item=True)
        elif tag == "h1":
            self._new_paragraph()
            self._current_paragraph.style = self.document.styles["Heading 1"]
        elif tag == "h2":
            self._new_paragraph()
            self._current_paragraph.style = self.document.styles["Heading 2"]
        elif tag == "br":
            self._current_paragraph.add_run().add_break()

    def handle_endtag(self, tag: str) -> None:
        if tag in {"strong", "b", "em", "i"}:
            self._current_runs = []
        elif tag in {"p", "div", "section", "li"}:
            self._current_paragraph = self.document.add_paragraph()
            self._current_paragraph.style = self.document.styles["Normal"]
        elif tag in {"ul", "ol"}:
            self._list_level = max(0, self._list_level - 1)

    def handle_data(self, data: str) -> None:
        text = data.replace("\xa0", " ")
        if not text.strip():
            return
        run = self._ensure_run()
        run.add_text(text)

    # Internal helpers -------------------------------------------------------
    def _ensure_run(self):
        if self._current_runs:
            return self._current_runs[-1]
        run = self._current_paragraph.add_run()
        run.font.size = Pt(11)
        if self._list_level:
            self._current_paragraph.style = self.document.styles["List Bullet"]
        self._current_runs.append(run)
        return run

    def _start_run(self, bold: bool = False, italic: bool = False) -> None:
        run = self._current_paragraph.add_run()
        run.font.size = Pt(11)
        run.bold = bold
        run.italic = italic
        self._current_runs.append(run)

    def _new_paragraph(self, list_item: bool = False) -> None:
        self._current_paragraph = self.document.add_paragraph()
        self._current_paragraph.style = (
            self.document.styles["List Bullet"] if list_item else self.document.styles["Normal"]
        )
        if list_item:
            self._current_paragraph.paragraph_format.left_indent = Pt(18)
        self._current_runs = []


def html_to_docx(html: str, output_path: Path, title: str | None = None) -> Path:
    document = Document()
    if title:
        head = document.add_heading(title, level=1)
        head.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    parser = _DocxHTMLParser(document)
    try:
        parser.feed(html)
    except Exception as exc:  # pragma: no cover - fallback logging
        LOGGER.warning("Simple HTML to DOCX parser failed: %s", exc)
        paragraph = document.add_paragraph()
        paragraph.add_run(html)
    document.core_properties.title = title or "Document"
    document.core_properties.language = "fr-FR"
    document.save(str(output_path))
    return output_path


__all__ = ["html_to_docx"]
