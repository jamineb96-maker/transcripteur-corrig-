"""Construction programmatique du gabarit DOCX de facture."""

from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path
from tempfile import NamedTemporaryFile
from typing import Dict, List

from docx import Document  # type: ignore
from docx.enum.style import WD_STYLE_TYPE  # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
from docx.oxml import OxmlElement  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.shared import Cm, Mm, Pt, RGBColor  # type: ignore

from .assets_bootstrap import AssetPaths

LOGGER = logging.getLogger(__name__)

DEFAULT_FONT = "Calibri"


def _format_currency(value: float) -> str:
    formatted = f"{value:,.2f}".replace(",", " ")
    return formatted.replace(".", ",")


def _ensure_style(document: Document, name: str, *, font_size: Pt, bold: bool = False, color: str | None = None,
                  align: WD_ALIGN_PARAGRAPH | None = None, space_after: Pt | None = None) -> None:
    styles = document.styles
    if name in styles:
        return
    style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = DEFAULT_FONT
    style.font.size = font_size
    style.font.bold = bold
    if color:
        rgb = RGBColor.from_string(color.replace("#", ""))
        style.font.color.rgb = rgb
    if space_after is not None:
        style.paragraph_format.space_after = space_after
    if align is not None:
        style.paragraph_format.alignment = align


def _create_styles(document: Document) -> None:
    _ensure_style(document, "TitleInvoice", font_size=Pt(16), bold=True, color="#111111", space_after=Pt(6))
    _ensure_style(document, "SmallGray", font_size=Pt(9), color="#555555")
    _ensure_style(document, "TableHeader", font_size=Pt(10), bold=True, color="#111111")
    _ensure_style(document, "TableCell", font_size=Pt(10))
    _ensure_style(document, "Totals", font_size=Pt(12), bold=True)
    _ensure_style(document, "Legal", font_size=Pt(9), color="#555555")


def _add_header(document: Document, invoice: dict, assets: AssetPaths) -> None:
    section = document.sections[0]
    header = section.header
    table = header.add_table(rows=1, cols=2, width=section.page_width)
    table.allow_autofit = True
    left_cell, right_cell = table.rows[0].cells

    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = right_para.add_run()
    try:
        run.add_picture(str(assets.logo_raster), width=Mm(34))
    except Exception as exc:  # pragma: no cover - dépend d'actifs
        LOGGER.warning("Insertion du logo impossible: %s", exc)
    info_lines: List[str] = []
    issuer = invoice.get("issuer", {})
    for key in ("name", "address", "email", "phone", "siret", "ape"):
        value = str(issuer.get(key, "")).strip()
        if value:
            info_lines.append(value)
    if info_lines:
        info_para = right_cell.add_paragraph("\n".join(info_lines), style="SmallGray")
        info_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _add_title(document: Document) -> None:
    paragraph = document.add_paragraph("FACTURE", style="TitleInvoice")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_invoice_info(document: Document, invoice: dict) -> None:
    table = document.add_table(rows=1, cols=2)
    table.autofit = True
    table.allow_autofit = True
    left_cell, right_cell = table.rows[0].cells

    patient = invoice.get("patient", {})
    patient_lines = ["Facturé à"]
    name = str(patient.get("name", "")).strip()
    address = str(patient.get("address", "")).strip()
    if name:
        patient_lines.append(name)
    if address:
        patient_lines.extend(address.splitlines())
    left_cell.text = "\n".join(patient_lines)

    info_lines = [
        f"Facture n° {invoice.get('number', invoice.get('id', ''))}",
        f"Date {invoice.get('date', '')}",
        "TVA non applicable, art. 293 B du CGI",
    ]
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_para.style = document.styles["SmallGray"]
    right_para.add_run("\n".join(info_lines))


def _add_lines_table(document: Document, invoice: dict) -> None:
    line_items = list(invoice.get("lines", []) or [])
    if not line_items:
        line_items = [
            {"label": "Séance psychopraticien 60 min", "qty": 1, "unit_price": 0.0, "vat_rate": 0},
        ]
    rows = 1 + len(line_items)
    table = document.add_table(rows=rows, cols=4)
    table.autofit = True
    try:
        table.style = document.styles["Table Grid"]
    except KeyError:
        pass

    headers = ["Prestation", "Qté", "PU", "Total"]
    for idx, label in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = label
        para = cell.paragraphs[0]
        para.style = document.styles["TableHeader"]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx else WD_ALIGN_PARAGRAPH.LEFT
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F0F0F0")
        shd.set(qn("w:val"), "clear")
        tc_pr.append(shd)

    for row_idx, entry in enumerate(line_items, start=1):
        cells = table.rows[row_idx].cells
        cells[0].text = str(entry.get("label", ""))
        qty = float(entry.get("qty", 0))
        unit = float(entry.get("unit_price", 0))
        total = float(entry.get("total", qty * unit))
        qty_str = f"{qty:.2f}".rstrip("0").rstrip(".")
        cells[1].text = qty_str.replace(".", ",")
        cells[2].text = _format_currency(unit)
        cells[3].text = _format_currency(total)
        for idx, cell in enumerate(cells[1:], start=1):
            para = cell.paragraphs[0]
            para.style = document.styles["TableCell"]
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cells[0].paragraphs[0].style = document.styles["TableCell"]

    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.get_or_add_tcBorders()
            for border_name in ("top", "bottom", "left", "right"):
                border = getattr(borders, border_name)
                border.val = "single"
                border.sz = 6
                border.color = "999999"


def _add_totals(document: Document, totals: dict) -> None:
    total_table = document.add_table(rows=1, cols=2)
    total_table.autofit = True
    row = total_table.rows[0]
    row.cells[0].text = "Total dû"
    row.cells[1].text = _format_currency(float(totals.get("total", totals.get("total_ttc", 0.0))))
    row.cells[0].paragraphs[0].style = document.styles["Totals"]
    amount_para = row.cells[1].paragraphs[0]
    amount_para.style = document.styles["Totals"]
    amount_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _add_legal_mentions(document: Document) -> None:
    now = datetime.now().strftime("%d/%m/%Y")
    mentions = [
        "Cette facture peut être transmise à votre mutuelle…",
        f"Document généré le {now}",
    ]
    for text in mentions:
        paragraph = document.add_paragraph(text, style="Legal")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_signature(document: Document, assets: AssetPaths) -> None:
    paragraph = document.add_paragraph("Pour acquit,")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    try:
        run = document.add_paragraph().add_run()
        run.add_picture(str(assets.signature), width=Mm(28))
        run.paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    except Exception as exc:  # pragma: no cover - dépend d'actifs
        LOGGER.warning("Insertion de la signature impossible: %s", exc)
    name_para = document.add_paragraph("Benjamin Tramoni")
    name_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    name_para.style = document.styles["TableCell"]


def build_invoice_docx(invoice: Dict, assets: AssetPaths) -> Path:
    """Construit un document DOCX conforme aux spécifications visuelles."""

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = document.styles["Normal"]
    styles.font.name = DEFAULT_FONT
    styles.font.size = Pt(10)

    _create_styles(document)
    _add_header(document, invoice, assets)
    _add_title(document)
    _add_invoice_info(document, invoice)
    _add_lines_table(document, invoice)
    _add_totals(document, invoice.get("totals", {}))
    _add_legal_mentions(document)
    _add_signature(document, assets)

    with NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        document.save(tmp.name)
        tmp_path = Path(tmp.name)
    LOGGER.info("build.docx ok pour %s", invoice.get("id", "temp"))
    return tmp_path


__all__ = ["build_invoice_docx"]
